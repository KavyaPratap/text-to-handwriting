import re
import docx
from docx.shared import Pt
from PIL import Image

with open("dummy.txt", "r") as txt_file:
    text = txt_file.read()

# Replace single spaces with multiple spaces
text = re.sub(r' ', '      ', text)

# Create a Word document
doc = docx.Document()

# Function to add an image of a letter to the document
def add_letter_image_to_doc(letter):
    try:
        image_filename = f"{ord(letter)}.png"  # Assumes image files are named after ASCII codes
        img = Image.open(image_filename)
        img_width, img_height = img.size
        run = paragraph.add_run()
        run.add_picture(image_filename, width=Pt(img_width), height=Pt(img_height))
    except Exception as e:
        # If an image is not found for the letter, add the letter as plain text
        run = paragraph.add_run(letter)

# Iterate through characters in the text
paragraph = None
for char in text:
    if char == "\n":
        doc.add_paragraph("")  # Add a new line
        paragraph = None  # Reset the current paragraph
    else:
        if paragraph is None:
            paragraph = doc.add_paragraph()  # Start a new paragraph if none exists
        add_letter_image_to_doc(char)

# Save the Word document
doc.save("output.docx")
